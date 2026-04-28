import streamlit as st
import docx
from docx.shared import RGBColor
import os
import io
import pythoncom
import win32com.client
from datetime import datetime, date
import tempfile

st.set_page_config(page_title="Gerador de Relatórios BOMPARC", layout="wide")

st.title("Gerador Automático de Relatório - BOMPARC")
st.markdown("Preencha o formulário abaixo para gerar o relatório de Inspeção Visual (Word e PDF).")

import pythoncom
import win32com.client
import os

# Função para converter DOCX para PDF de forma segura via COM nativo
def convert_to_pdf(docx_path, pdf_path):
    # O Streamlit roda em threads, o Windows exige CoInitialize() para comunicação COM
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 = PDF
        doc.Close()
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

# --- FORMULÁRIO ---
with st.form("relatorio_form"):
    st.subheader("1. Identificação e Local")
    col1, col2, col3 = st.columns(3)
    cliente = col1.text_input("Cliente", "DOF SUBSEA BRASIL SERVIÇOS LTDA")
    embarcacao = col2.text_input("Embarcação (Local do Teste)", "SKANDI CHIEFTAIN")
    endereco = col3.text_input("Endereço", "PORTO DO AÇU")

    col4, col5 = st.columns(2)
    certif = col4.text_input("CERTIF (Relatório Nº / Plaqueta)", "68430204-32")
    nf_os = col5.text_input("NF/ROMANEIO-OS (Ordem de Serviço)", "6843")
    
    st.subheader("2. Dados da Inspeção e Equipamento")
    col6, col7, col8 = st.columns(3)
    data_inspecao = col6.date_input("Data da Inspeção", date.today())
    ns = col7.text_input("NS (Número de Série)", "AST220S/03")
    item = col8.selectbox("Descrição do Material (Item)", ["Cinta Sling", "Cinta Circular", "Manilha", "Outro"])
    if item == "Outro":
        item = st.text_input("Especifique o Material")
        
    col9, col10, col11 = st.columns(3)
    capac = col9.number_input("Capacidade (Capac) em TON", min_value=0.1, value=3.0, step=0.1)
    dimensao = col10.text_input("Dimensão", "2 MTS")
    quantidade = col11.number_input("Quantidade", min_value=1, value=1, step=1)
    
    col12, col13 = st.columns(2)
    materia_prima = col12.text_input("Matéria Prima", "METAL")
    teste = col13.text_input("Teste", "N/A")
    
    st.subheader("3. Parâmetros e Parecer")
    col14, col15, col16 = st.columns(3)
    end = col14.text_input("END (Ensaio Não Destrutivo)", "N/A")
    aprov = col15.selectbox("APROV (Laudo Final)", ["APROVADO", "REPROVADO"])
    obs = col16.text_area("OBS (Observações)", "NÃO HOUVE / NONE")
    
    submit_button = st.form_submit_button("Gerar Relatórios")

if submit_button:
    with st.spinner("Processando documento..."):
        # Regras Inteligentes
        # 1. Data + 1 ano
        data_str = data_inspecao.strftime("%d/%m/%Y")
        try:
            data_validade_str = data_inspecao.replace(year=data_inspecao.year + 1).strftime("%d/%m/%Y")
        except ValueError: # Tratamento para ano bissexto (29/02)
            data_validade_str = data_inspecao.replace(year=data_inspecao.year + 1, day=28).strftime("%d/%m/%Y")
            
        # 2. Equipamento = [Item] + [Capac] + " TON x " + [Dimensao]
        equipamento = f"{item.upper()} {capac:g} TON X {dimensao.upper()}"
        
        # 3. Capac KG = Capac TON * 1000 com separador de milhar
        capac_kg = f"{int(capac * 1000):,} KG".replace(",", ".")
        
        # 4. Critério de Aceitação
        if item.upper() == "CINTA SLING":
            criterio = "ABNT NBR 15637-1"
        elif item.upper() == "CINTA CIRCULAR":
            criterio = "ABNT NBR 15637-2"
        else:
            criterio = "PADRÃO BOMPARC" # fallback
            
        # 5. OS
        ordem_servico = f"BPC-OSTC-{nf_os}"
        
        # Cores
        vermelho = RGBColor(255, 0, 0)
        
        # Dicionário de Substituições (Prefixo, Valor, Negrito, Cor, LimparProxParagrafo)
        replacements = [
            ("RELATÓRIO Nº (Number Report):", certif, True, None, False),
            ("CLIENTE (Client):", cliente, True, None, False),
            ("LOCAL DO TESTE (Local of the Test):", embarcacao, True, None, False),
            ("ENDEREÇO (Address):", endereco, True, None, False),
            ("EQUIPAMENTO (Equipment):", equipamento, True, None, False),
            ("SÉRIE (Serial Number):", ns, True, vermelho, False),
            ("DATA DA INSPEÇÃO (Date):", data_str, True, vermelho, False),
            ("DATA DE VALIDADE (Validity Date):", data_validade_str, True, vermelho, False),
            ("CRITÉRIO DE ACEITAÇÃO (Criteria of accept):", criterio, True, None, False),
            ("DIMENSÕES (Dimensiones):", dimensao, False, None, True),
            ("MATÉRIA PRIMA (Feedstock):", materia_prima, False, None, True),
            ("CARGA DE TRABALHO \n(Workload):", capac_kg, True, None, True),
            ("CARGA DE TRABALHO", capac_kg, True, None, True), # caso n tenha \n
            ("QUANTIDADE\n(Quantity):", f"{quantidade:02d} UNIDADE", True, None, True),
            ("QUANTIDADE", f"{quantidade:02d} UNIDADE", True, None, True),
            ("IDENTIFICAÇÃO DE PLAQUETA (ID Plate):", certif, True, None, False),
            ("ORDEM DE SERVIÇO (Service Order):", ordem_servico, True, None, False),
            ("RELATÓRIO DE ENSAIO NÃO DESTRUTIVO", end, True, None, False),
            ("LAUDO FINAL (Final Report):", aprov, True, None, False),
            ("Observações e Recomendações", obs, True, None, False)
        ]

        # Carregar o DOCX
        template_path = os.path.join(os.getcwd(), "template.docx")
        
        if not os.path.exists(template_path):
            st.error(f"Erro: Modelo 'template.docx' não encontrado na pasta {os.getcwd()}.")
            st.stop()
            
        doc = docx.Document(template_path)

        def replace_in_paragraphs(paragraphs):
            for i, p in enumerate(paragraphs):
                p_text = p.text
                for prefix, val, bold_val, color_val, clear_next in replacements:
                    if prefix in p_text:
                        style = None
                        if len(p.runs) > 0:
                            style = p.runs[0].style
                        
                        p.clear()
                        # Reconstroi o parágrafo
                        prefix_run = p.add_run(prefix + "  ")
                        if style: prefix_run.style = style
                        prefix_run.bold = True
                        
                        val_run = p.add_run(str(val))
                        if style: val_run.style = style
                        val_run.bold = bold_val
                        if color_val:
                            val_run.font.color.rgb = color_val
                            
                        # Limpar parágrafo seguinte se tiver vestígios de template (como em Matéria Prima)
                        if clear_next and i + 1 < len(paragraphs):
                            paragraphs[i+1].clear()
                            
                        break

                # Tratamento especial para as datas nas assinaturas (fim da tabela)
                if "02/04/2026" in p_text and not "DATA DA" in p_text:
                    p.clear()
                    new_run = p.add_run(data_str)
                    new_run.bold = False

        # Buscar nos parágrafos principais
        replace_in_paragraphs(doc.paragraphs)
        
        # Buscar nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

        # Salvar Docx
        output_docx = os.path.join(os.getcwd(), f"Relatorio_{certif}.docx")
        doc.save(output_docx)
        
        # Gerar PDF
        output_pdf = os.path.join(os.getcwd(), f"Relatorio_{certif}.pdf")
        convert_to_pdf(output_docx, output_pdf)
        
        st.success("Relatórios gerados com sucesso!")
        
        with open(output_docx, "rb") as d:
            st.download_button(
                label="📄 Baixar Relatório (Word)",
                data=d,
                file_name=f"Relatorio_{certif}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        with open(output_pdf, "rb") as p:
            st.download_button(
                label="📕 Baixar Relatório (PDF)",
                data=p,
                file_name=f"Relatorio_{certif}.pdf",
                mime="application/pdf"
            )
